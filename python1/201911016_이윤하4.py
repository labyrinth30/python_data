from docx import Document

# 100개의 다국어 지원 키워드
keywords = [
    "multilingual support", "language translation", "internationalization", "localization", "cross-cultural communication",
    "language diversity", "globalization", "translation services", "multilingual communication", "linguistic diversity",
    "language compatibility", "cultural adaptation", "language localization", "global audience", "translation tools",
    "multilingual content", "language services", "language technology", "cross-language communication", "language solutions",
    "multilingual website", "global communication", "multilingual applications", "language support", "translation platforms",
    "linguistic adaptation", "multilingual marketing", "language resources", "cross-cultural interaction", "language integration",
    "global language", "language accessibility", "multilingual interface", "cultural diversity", "language adaptation",
    "multilingual software", "language tools", "translation solutions", "multilingual data", "language processing",
    "cross-border communication", "multilingual user interface", "language globalization", "translation technology", "cultural compatibility",
    "language localization tools", "global communication tools", "multilingual resources", "language services platform", "cross-language understanding",
    "multilingual support system", "language adaptation tools", "translation management", "multilingual applications", "language diversity",
    "global communication solutions", "language translation tools", "cross-cultural communication tools", "multilingual platform", "language translation services",
    "international language support", "language translation platforms", "cross-cultural communication solutions", "multilingual marketing strategies", "language localization solutions",
    "globalization tools", "language support services", "cross-language communication solutions", "multilingual content management", "language translation technology",
    "global communication platforms", "language integration solutions", "cross-border communication tools", "multilingual support services", "language adaptation solutions",
    "global language support", "language accessibility solutions", "cross-language understanding tools", "multilingual user interface solutions", "language globalization platforms",
    "translation technology solutions", "multilingual resources management", "language services platform solutions", "cross-language understanding platforms", "multilingual support system solutions",
    "language adaptation tools", "translation management solutions", "multilingual applications development", "language diversity management", "global communication solutions development",
    "language translation tools development", "cross-cultural communication tools development", "multilingual platform development", "language translation services development", "international language support development"
]

# Word 문서 생성 및 키워드 쓰기
doc = Document()
doc.add_heading('Multilingual Support Keywords', level=1)

for keyword in keywords:
    doc.add_paragraph(keyword)

# Word 파일로 저장
doc.save('multilingual_support_keywords.docx')
